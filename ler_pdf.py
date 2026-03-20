import re
import csv
from collections import defaultdict
from pathlib import Path
from datetime import datetime
from urllib.parse import quote
import pdfplumber
from docx import Document

# ===== Config =====
ARQUIVO_PDF = "planos.pdf"
STATUSES = ["Pendente", "Prorrogado", "Vencido"]

# Saída por data de execução: saida/YYYY/MM/DD/...
hoje = datetime.now()
PASTA_SAIDA = Path("saida") / f"{hoje:%Y}" / f"{hoje:%m}" / f"{hoje:%d}"
PASTA_TXT = PASTA_SAIDA / "emails_txt"
PASTA_DOCX = PASTA_SAIDA / "emails_docx"
PASTA_REL = PASTA_SAIDA / "relatorios"

# Opcional: mapa líder -> email (para gerar link do Gmail com "Para:")
# Crie um arquivo "lideres_emails.csv" na pasta do projeto com colunas: lider,email
ARQUIVO_MAPA_EMAILS = Path("lideres_emails.csv")


# ===== Helpers =====
def limpar_nome_arquivo(nome: str) -> str:
    return re.sub(r'[<>:"/\\|?*]+', "", nome).strip()

def extrair_campo(texto: str, label: str) -> str:
    m = re.search(rf"^{re.escape(label)}\s*:\s*(.+)$", texto, re.MULTILINE)
    return m.group(1).strip() if m else ""

def extrair_linhas_planos(texto: str):
    padrao = re.compile(
        r"^(?P<linha>.+?)\s+(?P<data>\d{2}/\d{2}/\d{4})\s+(?P<status>Pendente|Vencido|Prorrogado)\s*$",
        re.MULTILINE | re.IGNORECASE,
    )
    planos = []
    for m in padrao.finditer(texto):
        planos.append({
            "data": m.group("data"),
            "status": m.group("status").capitalize(),
            "linha": m.group("linha").strip()
        })
    return planos

def texto_instrucao(status: str) -> str:
    if status == "Pendente":
        return ("Pedimos a gentileza de acompanhar o andamento do(s) plano(s) e assegurar o registro das ações realizadas no sistema, "
                "efetuando a atualização do status para REALIZADO ou PRORROGADO, conforme o caso, antes das respectivas datas limite.")
    if status == "Vencido":
        return ("Considerando o prazo já expirado, solicitamos a regularização no sistema, realizando a atualização do status para REALIZADO "
                "(caso a ação tenha sido concluída) ou PRORROGADO, conforme aplicável.")
    if status == "Prorrogado":
        return ("Ressaltamos que este é o prazo final após prorrogação, não sendo possível nova extensão.\n\n"
                "Solicitamos a verificação imediata e a atualização no sistema, alterando o status para REALIZADO ou NÃO CUMPRIDO, conforme a situação, até a data indicada.")
    return "Solicitamos a atualização do status no sistema conforme aplicável."

def data_mais_comum(datas: list[str]) -> str:
    freq = defaultdict(int)
    for d in datas:
        freq[d] += 1
    return sorted(freq.items(), key=lambda x: (-x[1], x[0]))[0][0] if datas else ""

def carregar_mapa_emails() -> dict:
    mapa = {}
    if not ARQUIVO_MAPA_EMAILS.exists():
        return mapa
    with ARQUIVO_MAPA_EMAILS.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        for row in reader:
            lider = (row.get("lider") or "").strip()
            email = (row.get("email") or "").strip()
            if lider and email:
                mapa[lider] = email
    return mapa

def gmail_compose_link(to: str, subject: str, body: str) -> str:
    # Link que abre o Gmail já com assunto/corpo (não salva rascunho automaticamente)
    base = "https://mail.google.com/mail/?view=cm&fs=1"
    params = []
    if to:
        params.append(f"to={quote(to)}")
    params.append(f"su={quote(subject)}")
    params.append(f"body={quote(body)}")
    return base + "&" + "&".join(params)

def montar_email_lider(lider: str, dados_lider: dict) -> tuple[str, str, dict]:
    contagens = {s: 0 for s in STATUSES}
    for st in dados_lider:
        for func in dados_lider[st]:
            contagens[st] += dados_lider[st][func]["qtd"]
    contagens["Total"] = sum(contagens.values())

    assunto = f"ALERTA DE PLANOS DE AÇÃO ({contagens['Total']} plano(s))"

    linhas = []
    linhas.append(f"Prezada {lider}, bom dia!")
    linhas.append("")

    for status in STATUSES:
        funcionarios = dados_lider.get(status, {})
        if not funcionarios:
            continue

        subtotal = sum(funcionarios[f]["qtd"] for f in funcionarios)
        linhas.append(f"Identificamos {subtotal} plano(s) de ação {status.upper()} sob sua responsabilidade, referente(s) ao(s) colaborador(es) abaixo:")
        linhas.append("")

        for funcionario in sorted(funcionarios.keys()):
            qtd = funcionarios[funcionario]["qtd"]
            data_ref = data_mais_comum(funcionarios[funcionario]["datas"])
            if qtd > 1:
                linhas.append(f"- {funcionario} — {qtd} plano(s) – Reavaliação prevista para {data_ref}")
            else:
                linhas.append(f"- {funcionario} – Reavaliação prevista para {data_ref}")

        linhas.append("")
        linhas.append(texto_instrucao(status))
        linhas.append("")
        linhas.append("Segue anexo o recorte do(s) plano(s) para conferência.")
        linhas.append("")
        linhas.append("Permanecemos à disposição.")
        linhas.append("")
        linhas.append("-" * 70)
        linhas.append("")

    corpo = "\n".join(linhas).rstrip()
    return assunto, corpo, contagens

def salvar_docx(caminho: Path, assunto: str, corpo: str):
    doc = Document()
    doc.add_heading("E-mail – Planos de Ação", level=1)
    doc.add_paragraph(f"Assunto: {assunto}")
    doc.add_paragraph("")

    # quebras de linha em parágrafos
    for linha in corpo.split("\n"):
        doc.add_paragraph(linha)

    doc.save(str(caminho))


# ===== Main =====
def main():
    # cria pastas de saída
    PASTA_TXT.mkdir(parents=True, exist_ok=True)
    PASTA_DOCX.mkdir(parents=True, exist_ok=True)
    PASTA_REL.mkdir(parents=True, exist_ok=True)

    mapa_emails = carregar_mapa_emails()

    # dados[lider][status][funcionario] = {"qtd": int, "datas": [...], "paginas": set()}
    dados = defaultdict(lambda: defaultdict(lambda: defaultdict(lambda: {"qtd": 0, "datas": [], "paginas": set()})))
    total_planos_pdf = 0

    with pdfplumber.open(ARQUIVO_PDF) as pdf:
        for pagina_idx, page in enumerate(pdf.pages, start=1):
            texto = page.extract_text() or ""
            funcionario = extrair_campo(texto, "Funcionário(a)")
            lider = extrair_campo(texto, "Superior Atual")

            if not funcionario or not lider:
                continue

            planos = extrair_linhas_planos(texto)
            for p in planos:
                status = p["status"]
                if status not in STATUSES:
                    continue

                total_planos_pdf += 1
                dados[lider][status][funcionario]["qtd"] += 1
                dados[lider][status][funcionario]["datas"].append(p["data"])
                dados[lider][status][funcionario]["paginas"].add(pagina_idx)

    # gera saída por líder
    lideres = sorted(dados.keys())

    # CSV de auditoria
    csv_path = PASTA_REL / "resumo.csv"
    with csv_path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["data_execucao", "lider", "pendente", "prorrogado", "vencido", "total", "gmail_link", "email_destino"])

        for lider in lideres:
            assunto, corpo, contagens = montar_email_lider(lider, dados[lider])

            nome_arq = limpar_nome_arquivo(lider)

            # TXT
            txt_path = PASTA_TXT / f"{nome_arq}.txt"
            txt_path.write_text(f"ASSUNTO: {assunto}\n\n{corpo}\n", encoding="utf-8")

            # DOCX
            docx_path = PASTA_DOCX / f"{nome_arq}.docx"
            salvar_docx(docx_path, assunto, corpo)

            # Link Gmail (opcional)
            email_destino = mapa_emails.get(lider, "")
            link = gmail_compose_link(email_destino, assunto, corpo)

            writer.writerow([
                f"{hoje:%Y-%m-%d}",
                lider,
                contagens["Pendente"],
                contagens["Prorrogado"],
                contagens["Vencido"],
                contagens["Total"],
                link,
                email_destino
            ])

    # resumo no terminal
    print(f"PDF: {ARQUIVO_PDF}")
    print(f"Data de execução: {hoje:%Y-%m-%d}")
    print(f"Total de líderes: {len(lideres)}")
    print(f"Total de planos (linhas) encontrados no PDF: {total_planos_pdf}")
    print(f"TXT gerados em: {PASTA_TXT.resolve()}")
    print(f"DOCX gerados em: {PASTA_DOCX.resolve()}")
    print(f"Resumo CSV: {csv_path.resolve()}")

if __name__ == "__main__":
    main()
